from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument

from fault_diagnosis_agent.models import FaultKnowledgeItem, KnowledgeStep
from fault_diagnosis_agent.retrieval.fault_types import classify_fault_type


class FaultDocumentParser:
    """Parse a Word operation manual into searchable fault knowledge items."""

    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)
        self.doc = DocxDocument(str(self.file_path))

    def parse(self) -> list[FaultKnowledgeItem]:
        items = self._parse_by_heading()
        if not items:
            items = self._parse_whole_document()
        return items

    def _parse_by_heading(self) -> list[FaultKnowledgeItem]:
        items: list[FaultKnowledgeItem] = []
        current_section = ""
        current_content: list[str] = []

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading"):
                if current_section:
                    items.extend(self._section_to_items(current_section, "\n".join(current_content)))
                current_section = text
                current_content = []
            else:
                current_content.append(text)

        if current_section:
            items.extend(self._section_to_items(current_section, "\n".join(current_content)))

        return items

    def _parse_whole_document(self) -> list[FaultKnowledgeItem]:
        text = "\n".join(para.text.strip() for para in self.doc.paragraphs if para.text.strip())
        return self._section_to_items(self.file_path.stem, text)

    def _section_to_items(self, section: str, content: str) -> list[FaultKnowledgeItem]:
        fault_type = classify_fault_type(f"{section}\n{content}")
        steps = extract_steps(content)
        risks = extract_risks(content)
        item = FaultKnowledgeItem(
            id=slugify(section),
            title=section,
            fault_type=fault_type,
            type=classify_section(section),
            aliases=[section],
            steps=steps,
            risks=risks,
            content=content,
            metadata={"source": str(self.file_path), "section": section},
        )
        return [item]


def classify_section(section_name: str) -> str:
    if "报警" in section_name or "处理流程" in section_name or "处置" in section_name:
        return "fault_procedure"
    if "逻辑" in section_name or "时机" in section_name:
        return "decision_logic"
    if "特征" in section_name or "变化" in section_name:
        return "phenomenon_pattern"
    return "general_knowledge"


def extract_steps(text: str) -> list[KnowledgeStep]:
    step_pattern = re.compile(r"(?:^|\n)\s*(\d+)[、.)．]\s*([^\n]+)")
    steps = [
        KnowledgeStep(step_id=f"S{match.group(1)}", text=match.group(2).strip(), source="文档步骤")
        for match in step_pattern.finditer(text)
    ]
    if steps:
        return steps

    candidates = re.split(r"[。；;]\s*", text)
    verbs = ("检查", "确认", "关闭", "开启", "观察", "记录", "排查", "联系", "调整")
    fallback = [line.strip() for line in candidates if any(verb in line for verb in verbs)]
    return [
        KnowledgeStep(step_id=f"S{i}", text=line, source="文档语句")
        for i, line in enumerate(fallback[:8], start=1)
        if line
    ]


def extract_risks(text: str) -> list[str]:
    risks = []
    for line in re.split(r"[\n。]", text):
        if any(word in line for word in ("风险", "注意", "防止", "避免", "异常说明")):
            cleaned = line.strip(" ：:")
            if cleaned:
                risks.append(cleaned)
    return risks[:5]


def slugify(value: str) -> str:
    value = re.sub(r"\s+", "-", value.strip())
    value = re.sub(r"[^\w\-\u4e00-\u9fff]", "", value)
    return value[:80] or "knowledge-item"


def save_knowledge(items: Iterable[FaultKnowledgeItem], output_path: str | Path) -> None:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    data = [item.model_dump() for item in items]
    output.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

